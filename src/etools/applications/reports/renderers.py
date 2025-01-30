from io import BytesIO

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from unicef_rest_export.renderers import ExportDocxTableRenderer

from etools.applications.reports.models import LowerResult


class ResultFrameworkRenderer(ExportDocxTableRenderer):
    def export_set(self, formatted, headers):
        stream = BytesIO()
        doc = Document()

        if not headers:
            doc.add_paragraph("No data provided.")
        else:
            table = doc.add_table(
                rows=1,
                cols=len(headers),
                style="Table Grid",
            )

            # set heading text
            header_cells = table.rows[0].cells
            for i, heading in enumerate(headers):
                header_cells[i].text = heading
                shader = parse_xml(
                    r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w'))
                )
                header_cells[i]._tc.get_or_add_tcPr().append(shader)

            # set data
            record_previous = None
            row_previous = None
            is_lowerresult = False
            for record in formatted:
                row = table.add_row().cells
                for i, key in enumerate(record):
                    if is_lowerresult or isinstance(record[key], LowerResult):
                        is_lowerresult = True
                        if key == "Result":
                            # Ensure record_previous is valid and has the key
                            if record_previous and key in record_previous:
                                result_previous = record_previous[key]
                            else:
                                result_previous = None

                            # Check both result_previous and record[key] before accessing .pk
                            if (
                                isinstance(result_previous, LowerResult) and
                                isinstance(record[key], LowerResult) and
                                result_previous is not None and
                                record[key] is not None
                            ):
                                if result_previous.pk == record[key].pk:
                                    # merge cells
                                    row_previous[i].merge(row[i])
                                    continue
                            row[i].text = str(record[key].name)
                        else:
                            row[i].text = str(record[key])
                    else:
                        is_lowerresult = False
                        if i < 2:
                            shader = parse_xml(
                                r'<w:shd {} w:fill="fef2cc"/>'.format(
                                    nsdecls('w')
                                )
                            )
                        else:
                            shader = parse_xml(
                                r'<w:shd {} w:fill="d9d9d9"/>'.format(
                                    nsdecls('w')
                                )
                            )
                        row[i]._tc.get_or_add_tcPr().append(shader)
                        row[i].text = str(record[key])
                    record_previous = record
                    row_previous = row

        doc.save(stream)
        return stream.getvalue()
